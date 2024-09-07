import os
from package.common_file import CommonFile
from docx import Document


class DocxFile(CommonFile):
    def __init__(self, out_dir, file_name):
        self.file_name = file_name
        super().__init__(out_dir)
        self.doc = Document()
        self.doc.add_heading(self.COMMON_FILE_HEADER, level=1)
    
    # 保存文件
    def add_paragraph(self, content):
        self.doc.add_paragraph(content)
    
    def add_heading(self, content, level):
        self.doc.add_heading(content, level=level)
    
    def save(self):
        # 保存 Word 文档
        dir_name = os.path.join(self.out_dir, self.file_name)
        self.doc.save(dir_name)
        print("Word document created: " + dir_name)